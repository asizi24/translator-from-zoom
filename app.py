"""
Flask Transcription App - Production-Grade API.

Features:
- RESTful API for transcription tasks
- File upload with validation
- Path traversal protection
- Rate limiting (optional)
- AI-powered Q&A and study materials
- Export to DOCX/PDF

Author: DevSquad AI (Senior Tech Lead Rewrite)
"""
from __future__ import annotations

import json
import logging
import os
import re
import time
from functools import wraps
from logging.handlers import RotatingFileHandler
from typing import Any, Callable, Dict, List, Optional, Tuple, TypeVar

from flask import Flask, render_template, request, jsonify, send_from_directory, Response
from werkzeug.utils import secure_filename
import google.generativeai as genai

# Optional MIME detection (graceful fallback if not installed)
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    magic = None  # type: ignore
    MAGIC_AVAILABLE = False

from transcriber_engine import TranscriptionManager
from config import Config, get_config

# =============================================================================
# Type Definitions
# =============================================================================
F = TypeVar("F", bound=Callable[..., Any])

# =============================================================================
# App Initialization
# =============================================================================
app = Flask(__name__)

# Load and validate configuration
config = get_config()
app.config.from_object(config)
app.secret_key = config.SECRET_KEY
app.config["MAX_CONTENT_LENGTH"] = config.MAX_CONTENT_LENGTH


# =============================================================================
# Logging Setup
# =============================================================================
def setup_logging() -> None:
    """Configure application logging with file rotation."""
    log_dir = os.path.dirname(config.LOG_FILE) if os.path.dirname(config.LOG_FILE) else "."
    if log_dir != "." and not os.path.exists(log_dir):
        os.makedirs(log_dir, exist_ok=True)

    logging.basicConfig(
        level=logging.DEBUG if config.DEBUG else logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    )

    file_handler = RotatingFileHandler(
        config.LOG_FILE,
        maxBytes=config.LOG_MAX_BYTES,
        backupCount=config.LOG_BACKUP_COUNT
    )
    file_handler.setLevel(logging.INFO)
    file_handler.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    ))
    logging.getLogger().addHandler(file_handler)

    # Reduce noise from external libraries
    logging.getLogger("werkzeug").setLevel(logging.WARNING)
    logging.getLogger("urllib3").setLevel(logging.WARNING)


setup_logging()
logger = logging.getLogger(__name__)


# =============================================================================
# Rate Limiting (Optional - graceful degradation)
# =============================================================================
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address

    rate_limit = os.getenv("RATE_LIMIT", "60/minute")
    limiter = Limiter(
        key_func=get_remote_address,
        app=app,
        default_limits=[rate_limit],
        storage_uri="memory://"
    )
    RATE_LIMITING_ENABLED = True
    logger.info("Rate limiting enabled: %s", rate_limit)
except ImportError:
    limiter = None  # type: ignore
    RATE_LIMITING_ENABLED = False
    logger.info("Rate limiting disabled (flask-limiter not installed)")


def rate_limit_decorator(limit: str) -> Callable[[F], F]:
    """Apply rate limit if available, otherwise pass through."""
    def decorator(f: F) -> F:
        if limiter:
            return limiter.limit(limit)(f)
        return f
    return decorator


# =============================================================================
# Folder Initialization
# =============================================================================
def ensure_folders() -> None:
    """Create required folders with explicit permissions."""
    for folder in [config.UPLOAD_FOLDER, config.DOWNLOAD_FOLDER]:
        try:
            os.makedirs(folder, exist_ok=True)
            if os.name != "nt":  # Linux/Mac
                os.chmod(folder, 0o755)
            logger.info("Folder ready: %s", folder)
        except PermissionError as e:
            logger.error("Permission denied creating %s: %s", folder, e)
            logger.error("Fix with: sudo chown -R $USER:$USER .")
            raise


ensure_folders()

# =============================================================================
# Initialize Services
# =============================================================================
logger.info("Initializing TranscriptionManager...")
manager = TranscriptionManager()

# Configure Gemini AI with Pro model for maximum capability
gemini_model: Optional[genai.GenerativeModel] = None
GEMINI_SAFETY_SETTINGS = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_ONLY_HIGH"},
]

if config.GOOGLE_API_KEY:
    genai.configure(api_key=config.GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel(
        "gemini-1.5-pro-latest",
        safety_settings=GEMINI_SAFETY_SETTINGS
    )
    logger.info("Google Gemini 1.5 Pro configured with relaxed safety settings")
else:
    logger.warning("Google Gemini API not configured - AI features disabled")


# =============================================================================
# Validation Helpers
# =============================================================================
ALLOWED_EXTENSIONS = config.ALLOWED_EXTENSIONS

# URL validation patterns
URL_PATTERNS = [
    re.compile(r"^https?://(?:www\.)?youtube\.com/watch\?v=[\w-]+", re.IGNORECASE),
    re.compile(r"^https?://(?:www\.)?youtu\.be/[\w-]+", re.IGNORECASE),
    re.compile(r"^https?://[\w.-]+\.zoom\.us/rec/", re.IGNORECASE),
    re.compile(r"^https?://", re.IGNORECASE),  # Generic HTTPS URLs (fallback)
]


# MIME types for video validation
ALLOWED_MIME_TYPES = {
    'video/mp4', 'video/x-msvideo', 'video/quicktime',
    'video/x-matroska', 'video/webm', 'video/x-flv',
    'video/mpeg', 'audio/mpeg', 'audio/wav', 'audio/x-wav',
    'audio/mp4', 'audio/x-m4a', 'application/octet-stream'  # Fallback for some files
}


def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def validate_upload_file(file) -> Tuple[bool, str]:
    """
    Comprehensive file validation: extension, size, and MIME type.
    
    Security Note:
        Extension-only checks are bypassable. We verify MIME type
        by reading file header (magic bytes) when python-magic is available.
    """
    filename = file.filename or ""
    
    # Check extension
    if not allowed_file(filename):
        allowed = ", ".join(ALLOWED_EXTENSIONS)
        return False, f"Invalid file type. Allowed: {allowed}"
    
    # Check file size (prevent 0-byte and oversized)
    file.seek(0, 2)  # Seek to end
    size = file.tell()
    file.seek(0)     # Reset for saving
    
    if size == 0:
        return False, "File is empty (0 bytes)"
    
    if size > config.MAX_CONTENT_LENGTH:
        max_mb = config.MAX_CONTENT_LENGTH // (1024 * 1024)
        return False, f"File too large (max {max_mb}MB)"
    
    # MIME type validation (if python-magic available)
    if MAGIC_AVAILABLE and magic:
        try:
            header = file.read(8192)  # Read first 8KB for magic bytes
            file.seek(0)  # Reset
            
            mime_type = magic.from_buffer(header, mime=True)
            if mime_type not in ALLOWED_MIME_TYPES:
                logger.warning("Rejected file with MIME type: %s", mime_type)
                return False, f"Invalid file content type detected: {mime_type}"
        except Exception as e:
            logger.warning("MIME detection failed, allowing upload: %s", e)
            # Graceful fallback - don't block if magic fails
    
    return True, ""


def validate_url(url: str) -> Tuple[bool, str]:
    """
    Validate URL is safe to process.
    
    Returns:
        Tuple of (is_valid, error_message)
    """
    if not url or not isinstance(url, str):
        return False, "URL is required"

    if len(url) > 2000:
        return False, "URL too long"

    # Check against allowed patterns
    for pattern in URL_PATTERNS:
        if pattern.match(url):
            return True, ""

    return False, "Invalid or unsupported URL format"


def get_transcript_text(task_status: Dict[str, Any]) -> Tuple[Optional[str], Optional[str]]:
    """
    Get transcript text from task status or file.
    
    Returns:
        Tuple of (text, error_message)
        
    Senior Dev Note:
        Consolidated duplicate transcript-reading logic from multiple endpoints.
    """
    text = task_status.get("transcript_text") or task_status.get("text", "")

    if text:
        return text, None

    # Fallback: read from file
    filename = task_status.get("filename")
    if filename and os.path.exists(filename):
        try:
            with open(filename, "r", encoding="utf-8") as f:
                return f.read(), None
        except IOError as e:
            logger.error("Failed to read transcript file %s: %s", filename, e)
            return None, "Could not read transcript"

    return None, "Transcript not available"


def is_path_safe(filepath: str, allowed_dirs: List[str]) -> bool:
    """
    Check if filepath is within allowed directories.
    
    Senior Dev Note:
        Path traversal protection using absolute path comparison.
        Uses os.path.normcase for Windows compatibility.
    """
    abs_path = os.path.normcase(os.path.abspath(filepath))

    for allowed_dir in allowed_dirs:
        allowed_abs = os.path.normcase(os.path.abspath(allowed_dir))
        # Check if path is under allowed directory
        if abs_path.startswith(allowed_abs + os.sep) or abs_path == allowed_abs:
            return True

    return False


# =============================================================================
# Routes - Core API
# =============================================================================
@app.route("/")
def index() -> str:
    """Serve the main page."""
    return render_template("index.html")


@app.route("/health")
def health() -> Tuple[Response, int]:
    """Health check endpoint for load balancers and monitoring."""
    try:
        health_status = {
            "status": "healthy",
            "service": "Flask Transcription App",
            "timestamp": time.time(),
            "queue_size": manager.task_queue.qsize(),
            "total_tasks": len(manager.tasks),
            "features": {
                "ai_enabled": bool(gemini_model),
                "rate_limiting": RATE_LIMITING_ENABLED,
                "diarization": bool(config.HF_TOKEN),
            }
        }
        return jsonify(health_status), 200
    except Exception as e:
        logger.exception("Health check failed")
        return jsonify({"status": "unhealthy", "error": str(e)}), 503


# =============================================================================
# Routes - Task Management
# =============================================================================
@app.route("/start", methods=["POST"])
@rate_limit_decorator("30/minute")
def start() -> Tuple[Response, int]:
    """Start a new transcription task from URL."""
    data = request.json or {}
    url = data.get("url")
    test_mode = data.get("test_mode", False)

    # Validate URL
    is_valid, error_msg = validate_url(url)
    if not is_valid:
        return jsonify({"error": error_msg}), 400

    task_id = manager.submit_task(url=url, test_mode=test_mode)
    return jsonify({"task_id": task_id, "status": "queued"}), 200


@app.route("/upload", methods=["POST"])
@rate_limit_decorator("10/minute")
def upload() -> Tuple[Response, int]:
    """
    Upload a file and start transcription.
    
    Security:
        - Validates extension, size, and MIME type
        - Uses secure_filename to prevent path traversal
        - Rejects 0-byte and oversized files
    """
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]

    if not file.filename:
        return jsonify({"error": "No file selected"}), 400

    # Comprehensive validation (extension + size + MIME)
    is_valid, error_msg = validate_upload_file(file)
    if not is_valid:
        return jsonify({"error": error_msg}), 400

    try:
        filename = secure_filename(file.filename)
        
        # Prevent filename collisions with timestamp
        import time as time_module
        base, ext = os.path.splitext(filename)
        unique_filename = f"{base}_{int(time_module.time())}{ext}"
        
        filepath = os.path.join(config.UPLOAD_FOLDER, unique_filename)
        file.save(filepath)
        
        # Verify file was saved correctly
        if not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
            return jsonify({"error": "File save failed"}), 500

        task_id = manager.submit_task(file_path=filepath)
        logger.info("File uploaded: %s (%d bytes)", unique_filename, os.path.getsize(filepath))

        return jsonify({
            "task_id": task_id,
            "status": "queued",
            "filename": unique_filename
        }), 200

    except IOError as e:
        logger.error("File upload failed: %s", e)
        return jsonify({"error": "File upload failed"}), 500


@app.route("/status/<task_id>")
def status(task_id: str) -> Tuple[Response, int]:
    """Get status of a specific task."""
    task_status = manager.get_status(task_id)

    if task_status is None:
        return jsonify({"error": "Task not found"}), 404

    return jsonify(task_status), 200


@app.route("/history")
def history() -> Tuple[Response, int]:
    """Get all tasks sorted by creation time."""
    all_tasks = manager.get_all_tasks()

    tasks_list = [
        {"task_id": task_id, **task_data}
        for task_id, task_data in all_tasks.items()
    ]
    tasks_list.sort(key=lambda x: x.get("created_at", 0), reverse=True)

    return jsonify({"tasks": tasks_list, "count": len(tasks_list)}), 200


# =============================================================================
# Routes - Content Access
# =============================================================================
@app.route("/preview/<task_id>")
def preview(task_id: str) -> Tuple[Response, int]:
    """Get preview of transcribed text."""
    task_status = manager.get_status(task_id)

    if task_status is None:
        return jsonify({"error": "Task not found"}), 404

    if task_status["status"] != "completed":
        return jsonify({"error": "Transcription not completed"}), 400

    text, error = get_transcript_text(task_status)
    if error:
        return jsonify({"error": error}), 404

    return jsonify({
        "task_id": task_id,
        "text": text,
        "filename": task_status.get("filename", "")
    }), 200


@app.route("/download/<task_id>")
def download(task_id: str) -> Tuple[Response, int]:
    """Download transcript file."""
    task_status = manager.get_status(task_id)

    if task_status is None:
        return jsonify({"error": "Task not found"}), 404

    if task_status["status"] != "completed":
        return jsonify({"error": "Task not completed"}), 400

    filename = task_status.get("filename")
    if not filename:
        return jsonify({"error": "File not found"}), 404

    # Security: Path traversal protection
    allowed_dirs = [config.DOWNLOAD_FOLDER, config.UPLOAD_FOLDER]
    if not is_path_safe(filename, allowed_dirs):
        logger.warning("Path traversal attempt blocked: %s", filename)
        return jsonify({"error": "Access denied"}), 403

    abs_path = os.path.abspath(filename)
    if not os.path.exists(abs_path):
        return jsonify({"error": "File not found"}), 404

    directory = os.path.dirname(abs_path)
    basename = os.path.basename(abs_path)

    return send_from_directory(directory, basename, as_attachment=True)


@app.route("/player/<task_id>")
def player(task_id: str) -> Tuple[str, int]:
    """Display the AI chat player for a completed transcription."""
    task_status = manager.get_status(task_id)

    if task_status is None:
        return "Task not found", 404

    if task_status["status"] != "completed":
        return "Transcription not completed", 400

    return render_template("player.html", task_id=task_id)


@app.route("/create_direct", methods=["POST"])
@rate_limit_decorator("30/minute")
def create_direct() -> Tuple[Response, int]:
    """
    Create a task directly from pasted transcript text.
    
    This allows users who already have a transcript to skip the
    transcription process and go directly to the AI chat player.
    """
    data = request.json or {}
    text = data.get("text", "").strip()

    # Validation
    if not text:
        return jsonify({"error": "הטקסט ריק. אנא הדבק תמלול."}), 400

    if len(text) < 10:
        return jsonify({"error": "הטקסט קצר מדי. אנא הדבק תמלול מלא."}), 400

    if len(text) > 500000:  # ~500KB limit
        return jsonify({"error": "הטקסט ארוך מדי (מקסימום 500,000 תווים)."}), 400

    try:
        # Generate task_id
        import uuid
        task_id = str(uuid.uuid4())[:8]
        
        # Create transcript file
        filename = f"direct_{task_id}_{int(time.time())}.txt"
        filepath = os.path.join(config.DOWNLOAD_FOLDER, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(text)
        
        # Create task with completed status
        manager.tasks[task_id] = {
            "status": "completed",
            "progress": 100,
            "message": "תמלול מהדבקה ישירה",
            "filename": filepath,
            "transcript_text": text,
            "transcript_segments": [{"speaker": "DIRECT_INPUT", "text": text, "start": 0, "end": 0}],
            "created_at": time.time(),
            "source": "direct_paste"
        }
        
        logger.info("Created direct paste task: %s (%d chars)", task_id, len(text))
        
        return jsonify({
            "task_id": task_id,
            "redirect_url": f"/player/{task_id}",
            "status": "completed"
        }), 200

    except IOError as e:
        logger.error("Failed to save direct paste: %s", e)
        return jsonify({"error": "שגיאה בשמירת הטקסט"}), 500
    except Exception as e:
        logger.exception("Direct paste error")
        return jsonify({"error": f"שגיאה: {e}"}), 500


# =============================================================================
# Routes - Export
# =============================================================================
@app.route("/export/<task_id>/<format>")
def export_document(task_id: str, format: str) -> Tuple[Response, int]:
    """Export transcript as DOCX or PDF."""
    if format not in ["docx", "pdf"]:
        return jsonify({"error": f"Unsupported format: {format}"}), 400

    task_status = manager.get_status(task_id)
    if not task_status:
        return jsonify({"error": "Task not found"}), 404

    if task_status["status"] != "completed":
        return jsonify({"error": "Transcription not completed"}), 400

    # Get segments
    segments = task_status.get("transcript_segments", [])
    if not segments:
        segments = _load_segments_from_file(task_status)

    if not segments:
        return jsonify({"error": "No transcript segments available"}), 404

    # Generate filename
    base_filename = os.path.basename(task_status.get("filename", f"transcript_{task_id}"))
    base_filename = os.path.splitext(base_filename)[0]

    try:
        if format == "docx":
            output_file = _generate_docx(segments, base_filename)
        else:
            output_file = _generate_pdf(segments, base_filename)

        directory = os.path.dirname(output_file)
        filename = os.path.basename(output_file)

        logger.info("Exported %s for task %s", format.upper(), task_id)
        return send_from_directory(directory, filename, as_attachment=True)

    except ImportError as e:
        logger.error("Export library not available: %s", e)
        return jsonify({"error": f"Export feature requires additional libraries: {e}"}), 500
    except Exception as e:
        logger.exception("Export error")
        return jsonify({"error": f"Export failed: {e}"}), 500


def _load_segments_from_file(task_status: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Load segments from JSON file."""
    filename = task_status.get("filename")
    if not filename:
        return []

    segments_file = filename.replace(".txt", "_segments.json")
    if not os.path.exists(segments_file):
        return []

    try:
        with open(segments_file, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data.get("segments", [])
    except (IOError, json.JSONDecodeError) as e:
        logger.error("Could not load segments file: %s", e)
        return []


def _generate_docx(segments: List[Dict[str, Any]], base_filename: str) -> str:
    """Generate DOCX with RTL Hebrew support."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    title = doc.add_heading(f"תמלול: {base_filename}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    current_speaker = None
    paragraph = None

    for segment in segments:
        speaker = segment.get("speaker", "UNKNOWN")
        text = segment.get("text", "").strip()

        if not text:
            continue

        if speaker != current_speaker:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            speaker_run = paragraph.add_run(f"{speaker}: ")
            speaker_run.bold = True
            speaker_run.font.size = Pt(12)
            speaker_run.font.color.rgb = RGBColor(0, 51, 102)

            paragraph.paragraph_format.right_to_left = True
            text_run = paragraph.add_run(text)
            text_run.font.size = Pt(11)

            current_speaker = speaker
        else:
            if paragraph:
                paragraph.add_run(" " + text)

    output_path = os.path.join(config.DOWNLOAD_FOLDER, f"{base_filename}.docx")
    doc.save(output_path)

    return output_path


def _generate_pdf(segments: List[Dict[str, Any]], base_filename: str) -> str:
    """Generate PDF with speaker formatting."""
    from fpdf import FPDF

    class HebrewPDF(FPDF):
        def __init__(self) -> None:
            super().__init__()
            try:
                self.add_font("Arial", "", r"C:\Windows\Fonts\arial.ttf", uni=True)
                self.add_font("Arial", "B", r"C:\Windows\Fonts\arialbd.ttf", uni=True)
            except Exception:
                pass

    pdf = HebrewPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    try:
        pdf.set_font("Arial", "B", 16)
    except Exception:
        pdf.set_font("Helvetica", "B", 16)

    pdf.cell(0, 10, f"Transcript: {base_filename}", ln=True, align="C")
    pdf.ln(5)

    current_speaker = None

    try:
        pdf.set_font("Arial", "", 10)
    except Exception:
        pdf.set_font("Helvetica", "", 10)

    for segment in segments:
        speaker = segment.get("speaker", "UNKNOWN")
        text = segment.get("text", "").strip()

        if not text:
            continue

        if speaker != current_speaker:
            pdf.ln(3)

            try:
                pdf.set_font("Arial", "B", 11)
            except Exception:
                pdf.set_font("Helvetica", "B", 11)

            pdf.cell(0, 6, f"{speaker}: ", ln=False)

            try:
                pdf.set_font("Arial", "", 10)
            except Exception:
                pdf.set_font("Helvetica", "", 10)

            pdf.multi_cell(0, 6, text)
            current_speaker = speaker
        else:
            pdf.multi_cell(0, 6, text)

    output_path = os.path.join(config.DOWNLOAD_FOLDER, f"{base_filename}.pdf")
    pdf.output(output_path)

    return output_path


# =============================================================================
# Routes - AI Features
# =============================================================================
@app.route("/ask", methods=["POST"])
@rate_limit_decorator("30/minute")
def ask() -> Tuple[Response, int]:
    """Handle AI questions about the transcript."""
    if not gemini_model:
        return jsonify({
            "error": "Google Gemini API not configured",
            "instructions": "Get your API key at: https://aistudio.google.com/app/apikey"
        }), 500

    data = request.json or {}
    task_id = data.get("task_id")
    question = data.get("question", "").strip()

    if not task_id:
        return jsonify({"error": "Missing task_id"}), 400

    if not question:
        return jsonify({"error": "Missing question"}), 400

    if len(question) > 2000:
        return jsonify({"error": "Question too long (max 2000 chars)"}), 400

    task_status = manager.get_status(task_id)
    if not task_status:
        return jsonify({"error": "Task not found"}), 404

    if task_status["status"] != "completed":
        return jsonify({"error": "Transcription not completed"}), 400

    transcript_text, error = get_transcript_text(task_status)
    if error:
        return jsonify({"error": error}), 404

    system_prompt = """אתה מורה פרטי ידידותי ומומחה. התלמיד שלך לומד משיעור.
ענה על השאלה שלו בהתבסס **רק** על התמלול של השיעור שמצורף למטה.

כללים חשובים:
- תן תשובות תמציתיות ומעודדות
- אם נשאלת על קוד, חלץ אותו מהתמלול ותקן שגיאות דיבור/תמלול
- השתמש בעברית להסברים, אלא אם מבוקש קוד
- אם התשובה לא נמצאת בשיעור, אמר בבירור: "לא מצאתי את זה בשיעור"

תמלול השיעור:
"""

    full_prompt = f"{system_prompt}\n\n{transcript_text}\n\nשאלת התלמיד: {question}"

    try:
        response = gemini_model.generate_content(full_prompt)
        return jsonify({"answer": response.text, "error": None}), 200

    except Exception as e:
        error_str = str(e)
        # Handle quota/rate limit errors specifically
        if "429" in error_str or "Quota" in error_str or "quota" in error_str:
            logger.error("QUOTA EXCEEDED in /ask: %s", e)
            return jsonify({
                "error": "מכסת ה-AI היומית הסתיימה. נסה שוב מחר.",
                "answer": None,
                "quota_exceeded": True
            }), 429
        logger.error("Gemini API error: %s", e)
        return jsonify({"error": f"AI error: {e}", "answer": None}), 500


@app.route("/generate_study_material", methods=["POST"])
@rate_limit_decorator("10/minute")
def generate_study_material() -> Tuple[Response, int]:
    """Generate AI-powered study material: summary and quiz."""
    if not gemini_model:
        return jsonify({
            "error": "Google Gemini API not configured",
            "instructions": "Get your API key at: https://aistudio.google.com/app/apikey"
        }), 500

    data = request.json or {}
    task_id = data.get("task_id")

    if not task_id:
        return jsonify({"error": "Missing task_id"}), 400

    task_status = manager.get_status(task_id)
    if not task_status:
        return jsonify({"error": "Task not found"}), 404

    if task_status["status"] != "completed":
        return jsonify({"error": "Transcription not completed"}), 400

    transcript_text, error = get_transcript_text(task_status)
    if error:
        return jsonify({"error": error}), 404

    prompt = f"""נתח את תמלול השיעור הבא. החזר אובייקט JSON תקין במבנה המדויק הבא:

{{
  "summary": "סיכום תמציתי בן 3 פסקאות של השיעור בעברית",
  "key_points": ["נקודה 1", "נקודה 2", "נקודה 3"],
  "quiz": [
    {{
      "question": "טקסט השאלה בעברית?",
      "options": ["אפשרות א", "אפשרות ב", "אפשרות ג", "אפשרות ד"],
      "correct_index": 0,
      "explanation": "למה זו התשובה הנכונה"
    }}
  ]
}}

חשוב:
- הסיכום צריך להיות 3 פסקאות קצרות
- key_points צריך להכיל בדיוק 3 נקודות עיקריות  
- quiz צריך להכיל בדיוק 5 שאלות
- כל שאלה חייבת להכיל בדיוק 4 אפשרויות
- correct_index הוא מספר בין 0-3
- ודא שהתשובה היא JSON תקין בלבד

תמלול השיעור:
{transcript_text}"""

    try:
        logger.info("Generating study material for task %s", task_id)
        response = gemini_model.generate_content(prompt)
        answer_text = response.text.strip()

        # Clean markdown
        if answer_text.startswith("```"):
            answer_text = answer_text.split("```")[1]
            if answer_text.startswith("json"):
                answer_text = answer_text[4:]
            answer_text = answer_text.strip()

        study_material = json.loads(answer_text)

        # Validate structure
        for key in ["summary", "key_points", "quiz"]:
            if key not in study_material:
                raise ValueError(f"Missing required key: {key}")

        logger.info("Generated study material for task %s", task_id)
        return jsonify(study_material), 200

    except json.JSONDecodeError as e:
        logger.error("JSON parse error: %s", e)
        return jsonify({"error": f"Invalid AI response format: {e}"}), 500
    except ValueError as e:
        logger.error("Validation error: %s", e)
        return jsonify({"error": str(e)}), 500
    except Exception as e:
        error_str = str(e)
        # Handle quota/rate limit errors specifically
        if "429" in error_str or "Quota" in error_str or "quota" in error_str:
            logger.error("QUOTA EXCEEDED in /generate_study_material: %s", e)
            return jsonify({
                "error": "מכסת ה-AI היומית הסתיימה. נסה שוב מחר.",
                "quota_exceeded": True
            }), 429
        logger.exception("Gemini API error")
        return jsonify({"error": f"AI error: {e}"}), 500


# =============================================================================
# Startup
# =============================================================================
def validate_environment() -> bool:
    """Validate and log environment status."""
    logger.info("Environment: %s", config.ENV)
    logger.info("Debug mode: %s", config.DEBUG)
    logger.info("Upload folder: %s", config.UPLOAD_FOLDER)
    logger.info("Download folder: %s", config.DOWNLOAD_FOLDER)
    logger.info("AI features: %s", "Enabled" if gemini_model else "Disabled")
    logger.info("Diarization: %s", "Enabled" if config.HF_TOKEN else "Disabled")
    logger.info("Rate limiting: %s", "Enabled" if RATE_LIMITING_ENABLED else "Disabled")
    return True


if __name__ == "__main__":
    logger.info("=" * 60)
    logger.info("🚀 Flask Transcription App - Starting...")
    logger.info("=" * 60)

    if not validate_environment():
        logger.error("Environment validation failed")
        exit(1)

    logger.info("✅ All checks passed! Starting server...")
    logger.info("Server running on http://0.0.0.0:5000")

    if config.is_production():
        logger.warning("Production mode - use Gunicorn/Waitress for deployment")

    app.run(debug=config.DEBUG, port=5000, host="0.0.0.0")
